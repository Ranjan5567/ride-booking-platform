"""
Convert SYSTEM-DESIGN-DOCUMENT.md to Word document
Requires: pip install python-docx markdown
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def markdown_to_word(md_file, output_file):
    """Convert markdown file to Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (but add spacing)
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            # H1
            p = doc.add_heading(line[2:], level=1)
            i += 1
        elif line.startswith('## '):
            # H2
            p = doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            # H3
            p = doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('#### '):
            # H4
            p = doc.add_heading(line[5:], level=4)
            i += 1
        elif line.startswith('---'):
            # Horizontal rule - add spacing
            doc.add_paragraph()
            i += 1
        elif line.startswith('```'):
            # Code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                # Make it look like code (monospace)
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
        elif line.startswith('|'):
            # Table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) > 1:
                # Parse table
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                    hdr_cells[j].paragraphs[0].runs[0].font.bold = True
                
                # Add data rows (skip separator line)
                for row_line in table_lines[2:]:
                    row_cells = table.add_row().cells
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    for j, cell in enumerate(cells):
                        row_cells[j].text = cell
        elif line.startswith('!'):
            # Image
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text, img_path = match.groups()
                # Try to add image if it exists
                img_full_path = Path(md_file).parent / img_path
                if img_full_path.exists():
                    try:
                        doc.add_picture(str(img_full_path), width=Inches(6))
                        p = doc.add_paragraph(alt_text)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.runs[0].font.italic = True
                        p.runs[0].font.size = Pt(9)
                    except:
                        # If image can't be added, just add the path as text
                        p = doc.add_paragraph(f"[Image: {alt_text} - {img_path}]")
                else:
                    p = doc.add_paragraph(f"[Image: {alt_text} - {img_path}]")
            i += 1
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet list
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
        elif re.match(r'^\d+\. ', line):
            # Numbered list
            p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            i += 1
        else:
            # Regular paragraph
            # Handle bold and italic
            text = line
            p = doc.add_paragraph()
            
            # Simple markdown parsing for bold
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            
            i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✓ Word document created: {output_file}")

if __name__ == '__main__':
    md_file = Path(__file__).parent / 'SYSTEM-DESIGN-DOCUMENT.md'
    output_file = Path(__file__).parent / 'SYSTEM-DESIGN-DOCUMENT.docx'
    
    print("Converting markdown to Word...")
    print("Note: This is a basic converter. For best results, use pandoc:")
    print("  pandoc SYSTEM-DESIGN-DOCUMENT.md -o SYSTEM-DESIGN-DOCUMENT.docx")
    print()
    
    try:
        markdown_to_word(md_file, output_file)
        print(f"\n✓ Success! Word document saved to: {output_file}")
    except ImportError:
        print("Error: python-docx not installed.")
        print("Install it with: pip install python-docx")
    except Exception as e:
        print(f"Error: {e}")
        print("\nAlternative: Use online converter or pandoc:")
        print("  https://www.markdowntoword.com/")
        print("  or install pandoc: https://pandoc.org/installing.html")

